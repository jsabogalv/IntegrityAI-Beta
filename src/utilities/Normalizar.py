import os
import sys
import pandas as pd
import pyarrow as pa
import pyarrow.parquet as pq
import zipfile
import pdfplumber
from docx import Document
from PIL import Image
import io
import json

def read_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def read_excel_file(file_path):
    df = pd.read_excel(file_path)
    return df.to_csv(index=False, encoding='utf-8')

def read_word_file(file_path):
    doc = Document(file_path)
    return '\n'.join([para.text for para in doc.paragraphs])

def read_image_file(file_path):
    with open(file_path, 'rb') as file:
        img = Image.open(file)
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format=img.format)
        return img_byte_arr.getvalue()

def read_json_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return json.dumps(json.load(file), ensure_ascii=False)

def read_html_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def read_pdf_file(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ''
    return text

def read_zip_file(file_path):
    content = ""
    with zipfile.ZipFile(file_path, 'r') as zip_ref:
        for file_name in zip_ref.namelist():
            with zip_ref.open(file_name) as file:
                ext = os.path.splitext(file_name)[1].lower()
                try:
                    if ext in ['.txt', '.csv']:
                        content += file.read().decode('utf-8', errors='ignore')
                    elif ext in ['.xls', '.xlsx']:
                        df = pd.read_excel(file)
                        content += df.to_csv(index=False, encoding='utf-8')
                    elif ext in ['.doc', '.docx']:
                        doc = Document(io.BytesIO(file.read()))
                        content += '\n'.join([para.text for para in doc.paragraphs])
                    elif ext == '.json':
                        content += json.dumps(json.load(io.BytesIO(file.read())), ensure_ascii=False)
                    elif ext in ['.html', '.htm']:
                        content += file.read().decode('utf-8', errors='ignore')
                    elif ext == '.pdf':
                        pdf_content = ""
                        with pdfplumber.open(io.BytesIO(file.read())) as pdf:
                            for page in pdf.pages:
                                pdf_content += page.extract_text() or ''
                        content += pdf_content
                except Exception as e:
                    content += f"\nFailed to read {file_name} inside zip: {str(e)}"
    return content

def read_files_to_arrow(file_paths):
    file_paths_list = []
    content_list = []
    log_list = []

    for file_path in file_paths:
        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext in ['.txt', '.csv']:
                content = read_text_file(file_path)
            elif ext in ['.xls', '.xlsx']:
                content = read_excel_file(file_path)
            elif ext in ['.doc', '.docx']:
                content = read_word_file(file_path)
            elif ext in ['.jpg', '.jpeg', '.png']:
                content = read_image_file(file_path)
            elif ext == '.json':
                content = read_json_file(file_path)
            elif ext in ['.html', '.htm']:
                content = read_html_file(file_path)
            elif ext == '.pdf':
                content = read_pdf_file(file_path)
            elif ext == '.zip':
                content = read_zip_file(file_path)
            else:
                raise ValueError("Unsupported file format")
            
            file_paths_list.append(file_path)
            content_list.append(content)
            log_list.append([file_path, os.path.getsize(file_path), "Success"])
        except Exception as e:
            log_list.append([file_path, os.path.getsize(file_path), f"Failed: {str(e)}"])
    
    table = pa.table([file_paths_list, content_list], names=['file_path', 'content'])
    return table, log_list

def save_arrow_file(table, save_path):
    pq.write_table(table, save_path)

def get_next_file_number(directory, prefix, extension):
    existing_files = [f for f in os.listdir(directory) if f.startswith(prefix) and f.endswith(extension)]
    if not existing_files:
        return 1
    numbers = [int(f.split('_')[0]) for f in existing_files]
    return max(numbers) + 1

def unificar_y_normalizar_archivos(input_paths, output_dir):
    os.makedirs(output_dir, exist_ok=True)
    table, log_list = read_files_to_arrow(input_paths)

    next_number = get_next_file_number(output_dir, '', 'parquet')
    save_path = os.path.join(output_dir, f'{next_number:02d}_normalized_data.parquet')
    save_arrow_file(table, save_path)

    log_path = os.path.join(output_dir, f'{next_number:02d}_LOG.txt')
    with open(log_path, 'w', encoding='utf-8') as log_file:
        log_file.write('File Path\tSize (bytes)\tStatus\n')
        for log_entry in log_list:
            log_file.write(f"{log_entry[0]}\t{log_entry[1]}\t{log_entry[2]}\n")

    print(f"Data has been normalized and saved to {save_path}")
    print(f"Log has been saved to {log_path}")

if __name__ == "__main__":
    if len(sys.argv) >= 3:
        input_paths = sys.argv[1:-1]
        output_dir = sys.argv[-1]
        unificar_y_normalizar_archivos(input_paths, output_dir)
    else:
        print("Uso: Normalizar.py <ruta_de_archivo1> <ruta_de_archivo2> ... <ruta_de_salida>")
